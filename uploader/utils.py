import os
import re
import tempfile
import zipfile
import shutil
import traceback
import time
import pythoncom

# 尝试导入处理不同格式文档的库
try:
    from docx import Document  # 处理.docx文件
except ImportError:
    print("警告: python-docx库未安装，.docx文件处理功能将不可用")

try:
    import docx2txt  # 处理.doc和.docx文件的备选方案
except ImportError:
    print("警告: docx2txt库未安装，将尝试使用其他方法")
    docx2txt = None

def extract_text_with_win32com(file_path, max_retries=3):
    """使用win32com优化提取Word文档内容，增加重试机制和错误处理"""
    text = ""
    word = None
    doc = None
    
    for retry in range(max_retries):
        try:
            # 在多线程环境中需要初始化COM
            import pythoncom
            pythoncom.CoInitialize()
            
            # 使用 win32com 自动化 Word。某些情况下 win32com 的 gen_py 缓存会损坏，
            # 导致类似 "module 'win32com.gen_py.xxx' has no attribute 'CLSIDToPackageMap'" 的错误。
            # 为提高鲁棒性，按以下顺序尝试：
            # 1) 优先使用 DispatchEx（更适合多线程，且通常能避免 gen_py 的一些问题）
            # 2) 若失败，尝试使用 gencache.EnsureDispatch（可能会触发生成/修复 gen_py）
            # 3) 若 EnsureDispatch 失败，尝试 gencache.Rebuild 再 EnsureDispatch
            import win32com.client
            try:
                # DispatchEx 在多线程或并发环境下更安全
                word = win32com.client.DispatchEx("Word.Application")
            except Exception as de:
                # DispatchEx 失败时，尝试通过 gencache 修复生成的缓存并使用 Dispatch
                try:
                    from win32com.client import gencache
                    try:
                        # EnsureDispatch 有时会修复缺失的 gen_py 模块
                        gencache.EnsureDispatch("Word.Application")
                    except Exception:
                        # 如果 EnsureDispatch 也失败，尝试 Rebuild 再 EnsureDispatch
                        try:
                            gencache.Rebuild()
                            gencache.EnsureDispatch("Word.Application")
                        except Exception:
                            # 如果仍然失败，让外层捕获并记录错误
                            raise
                    # 最终以普通 Dispatch 取得对象
                    word = win32com.client.Dispatch("Word.Application")
                except Exception as e_gencache:
                    # 最后尝试使用动态派生（dynamic.Dispatch）以绕开 gen_py 缓存问题
                    try:
                        from win32com.client import dynamic
                        word = dynamic.Dispatch("Word.Application")
                    except Exception as e_dynamic:
                        # 将原始错误与 gencache / dynamic 错误都记录并抛出，以便外层捕获
                        raise RuntimeError(f"win32com Dispatch 失败: DispatchEx error={de}; gencache error={e_gencache}; dynamic error={e_dynamic}")

            # 设置可见性和警告行为
            word.Visible = False
            word.DisplayAlerts = False
            
            # 使用绝对路径并确保文件存在
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                raise FileNotFoundError(f"文件不存在: {abs_path}")
            
            # 尝试打开文档
            if file_path.lower().endswith('.doc'):
                doc = word.Documents.Open(abs_path)
            else:
                doc = word.Documents.Open(abs_path)
            
            # 短暂延迟确保文档完全加载
            time.sleep(0.5)
            
            # 提取文本内容
            text = doc.Content.Text
            
            # 清理资源
            if doc:
                doc.Close(SaveChanges=False)
                doc = None
            if word:
                word.Quit()
                word = None
                
            break  # 成功提取，退出重试循环
            
        except Exception as e:
            # 清理资源
            if doc:
                try:
                    doc.Close(SaveChanges=False)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass
            word = None
            doc = None
            
            if retry < max_retries - 1:
                print(f"win32com提取尝试 {retry+1} 失败: {e}，将重试...")
                time.sleep(1)  # 重试前等待1秒
            else:
                print(f"win32com提取失败（所有尝试）: {e}")
    
    return text

def read_word_document(file_path):
    """读取Word文档内容，支持多种方法"""
    full_text = []
    
    # 根据文件类型选择不同的读取方法
    if file_path.lower().endswith('.docx') and 'Document' in globals():
        # 使用python-docx处理.docx文件
        doc = Document(file_path)
        print(f"使用python-docx处理.docx文件")
        
        print(f"文档包含 {len(doc.paragraphs)} 个段落")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                full_text.append(para.text)
                if i < 3:  # 只打印前3个段落内容作为示例
                    print(f"段落 {i+1} 内容预览: {para.text[:50]}...")
        
        # 读取表格内容
        print(f"文档包含 {len(doc.tables)} 个表格")
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append('\t'.join(row_text))
    
    elif docx2txt is not None:
        # 尝试使用docx2txt处理
        try:
            print(f"尝试使用docx2txt处理文件")
            text = docx2txt.process(file_path)
            if text.strip():
                full_text = text.split('\n')
                print(f"成功提取文本内容，约 {len(full_text)} 行")
                if len(full_text) > 0:
                    print(f"内容预览: {full_text[0][:100]}...")
        except Exception as e:
            print(f"docx2txt处理失败: {e}，尝试使用win32com")
    
    # 无论前面是否成功，都尝试使用win32com作为备选方案
    if not full_text:
        try:
            print("尝试使用win32com处理文档")
            text = extract_text_with_win32com(file_path)
            if text.strip():
                full_text = text.split('\n')
                print(f"成功提取文本内容，约 {len(full_text)} 行")
                if len(full_text) > 0:
                    print(f"内容预览: {full_text[0][:100]}...")
        except Exception as inner_e:
            print(f"win32com处理失败: {inner_e}")
            # 最后尝试一个简单的方法 - 如果是.docx，可以作为zip解压读取XML内容
            if file_path.lower().endswith('.docx'):
                print("尝试将.docx作为zip文件解压读取")
                try:
                    with zipfile.ZipFile(file_path, 'r') as doc_zip:
                        # 读取document.xml文件
                        with doc_zip.open('word/document.xml') as xml_file:
                            xml_content = xml_file.read().decode('utf-8')
                            # 简单移除XML标签
                            text = re.sub(r'<[^>]+>', '', xml_content)
                            text = re.sub(r'\s+', ' ', text)
                            full_text = text.split(' ')
                            print(f"成功提取部分内容")
                except Exception as inner_inner_e:
                    print(f"XML解析失败: {inner_inner_e}")
                    print("无法提取文档内容")
    
    return full_text



def extract_price_info(text, price_type, patterns):
    """从文本中提取价格信息"""
    for pattern in patterns:
        match = pattern.search(text)
        if match:
            price_str = match.group(1)
            print(f"{price_type}: {price_str}元")
            return float(price_str.replace(',', ''))
    print(f"未找到{price_type}")
    return 0.0

def extract_maintenance_fee(text):
    """提取维护费（含税）合计"""
    patterns = [
        re.compile(r'维护费（含税）合计[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'维护费合计[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'维护费（含税）[：:]*\s*([\d,]+\.?\d*)元')
    ]
    return extract_price_info(text, "维护费（含税）", patterns)

def extract_overall_total_price(text):
    """提取总体估算价格"""
    patterns = [
        re.compile(r'总体估算[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'总体估算价格[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'总体估算\s*([\d,]+\.?\d*)元'),
        re.compile(r'项目总体合计（含税）[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'总体估算（含税）共计[\s:：]*([\d,]+\.?\d*)元'),
        re.compile(r'和商务总体估算（含税）共计[\s:：]*([\d,]+\.?\d*)元')
    ]
    for pattern in patterns:
        match = pattern.search(text)
        if match:
            overall_total_price_str = match.group(1)
            print(f"总体估算价格: {overall_total_price_str}元")
            return float(overall_total_price_str.replace(',', ''))
    print("未在文本中找到明确的总体估算价格")
    return None

def extract_total_price(text):
    """提取总估算价格"""
    patterns = [
        re.compile(r'项目合计（含税）总估算\s*([\d,]+\.?\d*)元'),
        re.compile(r'总估算\s*([\d,]+\.?\d*)元'),
        re.compile(r'总估算[：:]*\s*([\d,]+\.?\d*)元')
    ]
    for pattern in patterns:
        match = pattern.search(text)
        if match:
            total_price_str = match.group(1)
            print(f"总估算价格: {total_price_str}元")
            return float(total_price_str.replace(',', ''))
    print("未在文本中找到明确的总估算价格")
    return None

def extract_broadband_maintenance_fee(text):
    """提取宽带维护费价格"""
    patterns = [
        re.compile(r'宽带维护费（含税）[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'宽带维护费（含税）[\s:：]*[^=]*=([\d,]+\.?\d*)元'),
        re.compile(r'宽带维护费（含税）[\s:：]*([\d,]+\.?\d*)元'),
        re.compile(r'宽带维护费（含税）合计[：:]*\s*([\d,]+\.?\d*)元')
    ]
    return extract_price_info(text, "宽带维护费（含税）", patterns)

def extract_broadband_service_fee(text):
    """提取宽带服务费价格"""
    patterns = [
        re.compile(r'宽带服务费（含税）[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'宽带服务费（含税）[\s:：]*[^=]*=([\d,]+\.?\d*)元'),
        re.compile(r'宽带服务费（含税）[\s:：]*([\d,]+\.?\d*)元')
    ]
    return extract_price_info(text, "宽带服务费（含税）", patterns)

def extract_terminal_fee(text):
    """提取终端费价格"""
    patterns = [
        re.compile(r'终端费（含税）[：:]*\s*([\d,]+\.?\d*)元'),
        re.compile(r'终端费（含税）[\s:：]*[^=]*=([\d,]+\.?\d*)元'),
        re.compile(r'终端费（含税）[\s:：]*([\d,]+\.?\d*)元')
    ]
    return extract_price_info(text, "终端费（含税）", patterns)

def extract_fiber_length(text):
    """提取光缆长度 - 更灵活的正则表达式"""
    # 尝试多种可能的光缆长度格式
    fiber_patterns = [
        re.compile(r'光缆\s*([\d.]+)\s*米'),  # 光缆 123 米
        re.compile(r'光缆([\d.]+)米'),         # 光缆123米
        re.compile(r'光缆长度\s*[：:]\s*([\d.]+)\s*米'),  # 光缆长度：123米
        re.compile(r'光缆长度\s*为\s*([\d.]+)\s*米'),  # 光缆长度为123米
        re.compile(r'光缆\s*长度\s*[：:]\s*([\d.]+)\s*米'),  # 光缆 长度：123米
        re.compile(r'光缆\s*约\s*([\d.]+)\s*米'),  # 光缆约123米
        re.compile(r'([\d.]+)\s*米\s*光缆'),  # 123米光缆
        # 添加更多可能的格式
        re.compile(r'光缆\s*总长度\s*[：:]\s*([\d.]+)\s*米'),  # 光缆总长度：123米
        re.compile(r'光缆\s*总长\s*[：:]\s*([\d.]+)\s*米'),    # 光缆总长：123米
        re.compile(r'光纤\s*([\d.]+)\s*米'),  # 光纤 123 米
        re.compile(r'光纤([\d.]+)米'),         # 光纤123米
        re.compile(r'光缆\s*铺设\s*([\d.]+)\s*米'),  # 光缆铺设 123 米
        re.compile(r'铺设\s*光缆\s*([\d.]+)\s*米')   # 铺设光缆 123 米
    ]
    
    # 首先在过滤后的文本中搜索
    print(f"\n=== 开始搜索光缆长度信息 ===")
    print(f"过滤后文本预览: {text[:200]}...")
    
    fiber_length = None
    for pattern in fiber_patterns:
        fiber_match = pattern.search(text)
        if fiber_match:
            fiber_length_str = fiber_match.group(1)
            fiber_length = float(fiber_length_str)
            print(f"在过滤后文本中找到光缆长度: {fiber_length}米 (使用模式: {pattern.pattern})")
            # 返回整数以符合显示要求
            return int(fiber_length)
    
    return fiber_length

def verify_calculation(info):
    """进行验算比较"""
    if info['doc_maintenance_total'] is not None:
        # 使用近似比较，因为浮点数计算可能有精度问题
        if abs(info['total_fees'] - info['doc_maintenance_total']) < 0.0001:
            print("✓ 验算通过：计算总和与文档中维护费合计一致")
            return True
        else:
            print(f"✗ 验算失败：计算总和({info['total_fees']:.4f}元)与文档中维护费合计({info['doc_maintenance_total']:.4f}元)不一致")
            return False
    return False

def debug_keyword_search(text):
    """添加一些关键字搜索来帮助调试"""
    keywords = ['光缆', '光纤', '米', '长度']
    print("文本中包含的关键字:")
    for keyword in keywords:
        if keyword in text:
            print(f"  - '{keyword}' 在文本中找到")
        else:
            print(f"  - '{keyword}' 未找到")

def extract_info_from_word(file_path, original_name=None):
    """从单个Word文档中提取信息
    
    Args:
        file_path (str): 文件的物理路径
        original_name (str, optional): 原始上传的文件名，用于提取单号
    """
    print(f"开始处理Word文档: {file_path}")
    results = []
    file_name = original_name or os.path.basename(file_path)
    print(f"文件名: {file_name}")
    
    # 提取文件名中的英文数字代码部分
    match = re.match(r'([A-Za-z0-9_]+)', file_name)
    if match:
        code_part = match.group(1)
        print(f"提取到代码部分: {code_part}")
        
        # 读取Word文档内容
        try:
            print("开始读取Word文档内容...")
            full_text = read_word_document(file_path)
            
            # 不再过滤文本，直接使用完整的原始文本
            filtered_text = '\n'.join(full_text)
            marker_found = True
            
            if marker_found:
                # 提取各类价格信息
                info = {
                    'code_part': code_part,
                    'maintenance_fee': 0.0,
                    'service_fee': 0.0,
                    'terminal_fee': 0.0,
                    'total_fees': 0.0,
                    'doc_maintenance_total': None,
                    'overall_total_price': None,
                    'total_price': None,
                    'fiber_length': None,
                    'verification_passed': False,
                    'file_name': file_name,
                    'extraction_status': '成功'
                }
                
                print(f"\n=== 提取到的价格信息 ===")
                print(f"单号: {code_part}")
                
                # 提取维护费（含税）合计
                info['doc_maintenance_total'] = extract_maintenance_fee(filtered_text)
                
                # 提取总体估算价格
                info['overall_total_price'] = extract_overall_total_price(filtered_text)
                
                # 提取总估算价格
                info['total_price'] = extract_total_price(filtered_text)
                
                # 提取宽带维护费价格
                info['maintenance_fee'] = extract_broadband_maintenance_fee(filtered_text)
                
                # 提取宽带服务费价格
                info['service_fee'] = extract_broadband_service_fee(filtered_text)
                
                # 提取终端费价格
                info['terminal_fee'] = extract_terminal_fee(filtered_text)
                
                # 计算费用总和
                info['total_fees'] = info['maintenance_fee'] + info['service_fee'] + info['terminal_fee']
                print(f"宽带维护费、宽带服务费和终端费的总和: {info['total_fees']:.4f}元")
                
                # 提取光缆长度
                info['fiber_length'] = extract_fiber_length(filtered_text)
                
                if info['fiber_length'] is None:
                    print("未找到光缆长度信息")
                    debug_keyword_search(filtered_text)
                print("========================\n")
                
                # 进行验算比较
                info['verification_passed'] = verify_calculation(info)
                
                print(f"====================\n")
                results.append(info)
            else:
                print(f"未找到标记文本，跳过后续处理")
                results.append({
                    'code_part': code_part,
                    'file_name': file_name,
                    'fiber_length': None,
                    'extraction_status': '失败',
                    'error': '未找到关键标记文本'
                })
        except Exception as e:
            print(f"处理文件 {file_path} 时出错: {e}")
            print(traceback.format_exc())
            results.append({
                'code_part': code_part if 'code_part' in locals() else '未知',
                'file_name': file_name,
                'fiber_length': None,
                'extraction_status': '失败',
                'error': str(e)
            })
    else:
        print(f"无法从文件名 {file_name} 中提取代码部分")
        print(f"文件名格式: {file_name}")
        
    return results

def extract_info_from_zip(zip_path, original_name=None):
    """从ZIP文件中提取Word文档内容并解析价格信息

    Args:
        zip_path (str): ZIP文件的物理路径
        original_name (str, optional): 原始上传的ZIP文件名，用于提取单号
    """
    print(f"开始处理压缩文件: {zip_path}")
    results = []
    temp_dir = tempfile.mkdtemp()
    extracted_files = []
    
    # 从ZIP文件名提取单号，优先使用原始文件名
    zip_file_name = original_name or os.path.basename(zip_path)
    # 使用正则表达式提取单号部分 (例如从 EOSC_4712508269337893_KC... 中提取 EOSC_4712508269337893)
    zip_code_match = re.match(r'([A-Za-z0-9_]+)', zip_file_name)
    zip_code_part = zip_code_match.group(1) if zip_code_match else None
    if zip_code_part:
        print(f"从ZIP文件名提取到单号: {zip_code_part}")
    else:
        print(f"无法从ZIP文件名 {zip_file_name} 中提取单号")
    
    try:
        print(f"创建临时目录: {temp_dir}")
        
        # 首先验证文件是否为有效的ZIP文件
        if not zipfile.is_zipfile(zip_path):
            raise ValueError(f"提供的文件不是有效的ZIP文件: {zip_path}")
        
        # 解压zip文件
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            print(f"开始解压文件...")
            zip_ref.extractall(temp_dir)
            print(f"解压完成，解压到临时目录: {temp_dir}")
            
            # 获取解压后的所有文件
            print("开始获取解压后的文件列表...")
            for root, dirs, files in os.walk(temp_dir):
                print(f"在目录 {root} 中发现 {len(files)} 个文件")
                for file in files:
                    file_path = os.path.join(root, file)
                    extracted_files.append(file_path)
                    print(f"找到文件: {file_path}")
        
        print(f"总共解压出 {len(extracted_files)} 个文件")
        
        # 处理每个Word文档，包括.doc和.docx格式
        word_files = [f for f in extracted_files if f.lower().endswith(('.doc', '.docx')) and not f.lower().endswith('.bak')]
        print(f"找到 {len(word_files)} 个Word文件(.doc或.docx)")
        for file in word_files:
            print(f"- {os.path.basename(file)}")
        
        for file_path in word_files:
            print(f"\n处理Word文档: {file_path}")
            file_name = os.path.basename(file_path)
            print(f"文件名: {file_name}")
            
            # 优先使用来自原始ZIP文件名的单号，或从Word文档文件名中提取
            match = re.search(r'(EOSC_[A-Za-z0-9_\-]+)(?:[^A-Za-z0-9_\-]|$)', file_name)
            if match:
                code_part = match.group(1)
                print(f"从文件名提取到单号: {code_part}")
            else:
                print(f"无法从文件名 {file_name} 中提取单号")
                continue
                
            # 读取Word文档内容
            try:
                print("开始读取Word文档内容...")
                full_text = read_word_document(file_path)
                
                # 不再过滤文本，直接使用完整的原始文本
                filtered_text = '\n'.join(full_text)
                marker_found = True
                
                if marker_found:
                    # 提取各类价格信息
                    info = {
                        'code_part': code_part,
                        'maintenance_fee': 0.0,
                        'service_fee': 0.0,
                        'terminal_fee': 0.0,
                        'total_fees': 0.0,
                        'doc_maintenance_total': None,
                        'overall_total_price': None,
                        'total_price': None,
                        'fiber_length': None,
                        'verification_passed': False,
                        'file_name': file_name,
                        'extraction_status': '成功'
                    }
                    
                    print(f"\n=== 提取到的价格信息 ===")
                    print(f"单号: {code_part}")
                    
                    # 提取维护费（含税）合计
                    info['doc_maintenance_total'] = extract_maintenance_fee(filtered_text)
                    
                    # 提取总体估算价格
                    info['overall_total_price'] = extract_overall_total_price(filtered_text)
                    
                    # 提取总估算价格
                    info['total_price'] = extract_total_price(filtered_text)
                    
                    # 提取宽带维护费价格
                    info['maintenance_fee'] = extract_broadband_maintenance_fee(filtered_text)
                    
                    # 提取宽带服务费价格
                    info['service_fee'] = extract_broadband_service_fee(filtered_text)
                    
                    # 提取终端费价格
                    info['terminal_fee'] = extract_terminal_fee(filtered_text)
                    
                    # 计算费用总和
                    info['total_fees'] = info['maintenance_fee'] + info['service_fee'] + info['terminal_fee']
                    print(f"宽带维护费、宽带服务费和终端费的总和: {info['total_fees']:.4f}元")
                    
                    # 提取光缆长度
                    info['fiber_length'] = extract_fiber_length(filtered_text)
                    
                    if info['fiber_length'] is None:
                        print("未找到光缆长度信息")
                        debug_keyword_search(filtered_text)
                    print("========================\n")
                    
                    # 进行验算比较
                    info['verification_passed'] = verify_calculation(info)
                    
                    print(f"====================\n")
                    results.append(info)
                else:
                    print(f"未找到标记文本，跳过后续处理")
                    results.append({
                        'code_part': code_part,
                        'file_name': file_name,
                        'fiber_length': None,
                        'extraction_status': '失败',
                        'error': '未找到关键标记文本'
                    })
            except Exception as e:
                print(f"处理文件 {file_path} 时出错: {e}")
                print(traceback.format_exc())
                results.append({
                    'code_part': code_part if 'code_part' in locals() else '未知',
                    'file_name': file_name,
                    'fiber_length': None,
                    'extraction_status': '失败',
                    'error': str(e)
                })
    except Exception as e:
        print(f"处理压缩文件 {zip_path} 时出错: {e}")
        print(traceback.format_exc())
        # 添加错误信息到结果中，以便前端展示
        results.append({
            'error': str(e),
            'fiber_length': None,
            'extraction_status': '失败',
            'file_name': os.path.basename(zip_path)
        })

    finally:
        # 清理临时目录
        print(f"清理临时目录: {temp_dir}")
        shutil.rmtree(temp_dir)

    return results